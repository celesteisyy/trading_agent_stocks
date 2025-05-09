import pandas as pd
import numpy as np
import os
import logging
import matplotlib.pyplot as plt
import io
from docx import Document
from docx.shared import Inches
import gradio as gr
from typing import Dict, List, Tuple, Optional, Union

class PortfolioManager:
    """
    PortfolioManager handles:
    1. Portfolio simulation based on trading signals
    2. Performance metrics calculation
    3. Report generation
    4. Dashboard visualization
    5. Sentiment contribution analysis
    """
    
    def __init__(
        self,
        initial_capital: float = 100000.0,
        freq: str = 'B',
        transaction_cost: float = 0.001,
        output_dir: str = 'output'
    ):
        """
        Initialize portfolio manager.
        
        Args:
            initial_capital: Starting portfolio value
            freq: Data frequency ('B' for business days, 'D' for calendar days)
            transaction_cost: Cost per trade as fraction of trade value
            output_dir: Directory for saving reports
        """
        self.logger = logging.getLogger(__name__)
        self.initial_capital = initial_capital
        self.freq = freq
        self.transaction_cost = transaction_cost
        self.output_dir = output_dir
        
        # Set annualization factor based on frequency
        self.annual_factor = 252 if freq == 'B' else 365
        
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
    
    def backtest(self, orders_df: pd.DataFrame) -> pd.DataFrame:
        """
        Simulate portfolio performance from trading signals.
        
        Args:
            orders_df: DataFrame with index dates and columns ['order', 'price']
                       where order=1 buys (enter), -1 sells (exit).
        
        Returns:
            DataFrame with portfolio performance metrics
        """
        self.logger.info("Running portfolio backtest...")
        df = orders_df.copy().sort_index()
        
        # Ensure we have required columns
        if 'order' not in df.columns or 'price' not in df.columns:
            raise ValueError("orders_df must have 'order' and 'price' columns")
        
        # Calculate daily price returns
        df['price_return'] = df['price'].pct_change().fillna(0)
        
        # Build position series: 1 after buy, 0 after sell
        # Map orders: buy->1, sell->0, no order->NaN
        pos = df['order'].replace({1: 1, -1: 0})
        df['position'] = pos.ffill().fillna(0)
        
        # Number of trades (changes in position)
        df['trade'] = df['position'].diff().fillna(0).abs()
        df['trade_cost'] = df['trade'] * df['price'] * self.transaction_cost
        
        # Strategy returns
        df['returns'] = df['position'].shift(1).fillna(0) * df['price_return'] - df['trade_cost'] / self.initial_capital
        
        # Equity curve
        df['equity'] = (1 + df['returns']).cumprod() * self.initial_capital
        
        # Drawdown calculation
        df['peak'] = df['equity'].cummax()
        df['drawdown'] = (df['equity'] / df['peak'] - 1) * 100  # as percentage
        
        self.logger.info(f"Backtest complete. Final equity: ${df['equity'].iloc[-1]:.2f}")
        return df
    
    def compute_performance(self, portfolio_df: pd.DataFrame) -> Dict[str, float]:
        """
        Calculate performance metrics from portfolio results.
        
        Args:
            portfolio_df: Output from backtest() method
            
        Returns:
            Dictionary of performance metrics
        """
        # Extract required data
        equity = portfolio_df['equity']
        returns = portfolio_df['returns']
        
        # Total return
        total_return = equity.iloc[-1] / self.initial_capital - 1
        
        # Annualized return
        n_periods = len(returns)
        years = n_periods / self.annual_factor
        ann_return = (1 + total_return) ** (1 / years) - 1 if years > 0 else 0
        
        # Volatility (annualized)
        volatility = returns.std() * np.sqrt(self.annual_factor)
        
        # Sharpe ratio (assume risk-free ~0)
        sharpe = ann_return / volatility if volatility != 0 else 0
        
        # Maximum drawdown
        max_drawdown = portfolio_df['drawdown'].min()
        
        # Calmar ratio
        calmar = ann_return / abs(max_drawdown / 100) if max_drawdown != 0 else 0
        
        # Win rate
        daily_wins = (returns > 0).sum()
        win_rate = daily_wins / len(returns) if len(returns) > 0 else 0
        
        # Number of trades
        n_trades = portfolio_df['trade'].sum() / 2  # divide by 2 since each round-trip is 2 position changes
        
        # Average trade return
        trade_returns = []
        in_trade = False
        entry_price = 0
        
        for i, row in portfolio_df.iterrows():
            if row['trade'] == 1 and not in_trade:  # Enter trade
                entry_price = row['price']
                in_trade = True
            elif row['trade'] == 1 and in_trade:  # Exit trade
                if entry_price > 0:
                    trade_return = (row['price'] / entry_price) - 1
                    trade_returns.append(trade_return)
                in_trade = False
        
        avg_trade = np.mean(trade_returns) if trade_returns else 0
        
        # Return all metrics
        metrics = {
            'total_return': total_return,
            'ann_return': ann_return,
            'volatility': volatility,
            'sharpe_ratio': sharpe,
            'max_drawdown': max_drawdown / 100,  # Convert back to decimal
            'calmar_ratio': calmar,
            'win_rate': win_rate,
            'n_trades': n_trades,
            'avg_trade': avg_trade
        }
        
        self.logger.info(f"Performance metrics calculated: Total Return: {total_return:.2%}, Sharpe: {sharpe:.2f}")
        return metrics
    
    def analyze_sentiment_contribution(
        self,
        base_portfolio: pd.DataFrame,
        full_portfolio: pd.DataFrame,
        sentiment_df: pd.DataFrame
    ) -> Dict:
        """
        Analyze and quantify the contribution of sentiment analysis to strategy performance.
        
        Args:
            base_portfolio: Portfolio performance without sentiment
            full_portfolio: Portfolio performance with sentiment
            sentiment_df: Sentiment data used for analysis
            
        Returns:
            Dictionary with sentiment contribution metrics
        """
        self.logger.info("Analyzing sentiment contribution to performance...")
        
        # Calculate base metrics
        base_metrics = self.compute_performance(base_portfolio)
        
        # Calculate full metrics
        full_metrics = self.compute_performance(full_portfolio)
        
        # Calculate absolute and relative performance differences
        performance_diff = {
            'return_diff': full_metrics['total_return'] - base_metrics['total_return'],
            'return_diff_pct': (full_metrics['total_return'] - base_metrics['total_return']) / abs(base_metrics['total_return']) if base_metrics['total_return'] != 0 else float('inf'),
            'sharpe_diff': full_metrics['sharpe_ratio'] - base_metrics['sharpe_ratio'],
            'drawdown_diff': base_metrics['max_drawdown'] - full_metrics['max_drawdown'],
            'win_rate_diff': full_metrics['win_rate'] - base_metrics['win_rate']
        }
        
        # Identify sentiment-driven trades
        sentiment_driven_trades = []
        
        # Get dates where base and full strategies diverge
        base_orders = base_portfolio[base_portfolio['order'] != 0]['order']
        full_orders = full_portfolio[full_portfolio['order'] != 0]['order']
        
        # Convert to dictionaries for faster lookup
        base_orders_dict = base_orders.to_dict()
        full_orders_dict = full_orders.to_dict()
        
        # Find dates in full_orders not in base_orders or with different signals
        for date, order in full_orders_dict.items():
            if date not in base_orders_dict or base_orders_dict[date] != order:
                # This trade was influenced by sentiment
                # Get the sentiment value on this date
                if date in sentiment_df.index:
                    sentiment_val = sentiment_df.loc[date, 'avg_compound'] if 'avg_compound' in sentiment_df.columns else None
                    sentiment_driven_trades.append({
                        'date': date,
                        'order': order,
                        'price': full_portfolio.loc[date, 'price'],
                        'sentiment': sentiment_val
                    })
        
        # Calculate market condition correlations
        # When is sentiment most useful?
        
        # Split the backtest period into equal segments
        segments = 4
        period_length = len(full_portfolio) // segments
        period_returns = []
        
        for i in range(segments):
            start_idx = i * period_length
            end_idx = (i + 1) * period_length if i < segments - 1 else len(full_portfolio)
            
            # Calculate returns for this period
            base_period_return = base_portfolio['equity'].iloc[end_idx-1] / base_portfolio['equity'].iloc[start_idx] - 1
            full_period_return = full_portfolio['equity'].iloc[end_idx-1] / full_portfolio['equity'].iloc[start_idx] - 1
            
            # Calculate average sentiment for this period
            period_dates = full_portfolio.index[start_idx:end_idx]
            avg_sentiment = sentiment_df.loc[sentiment_df.index.intersection(period_dates), 'avg_compound'].mean() if 'avg_compound' in sentiment_df.columns else None
            sentiment_volatility = sentiment_df.loc[sentiment_df.index.intersection(period_dates), 'avg_compound'].std() if 'avg_compound' in sentiment_df.columns else None
            
            period_returns.append({
                'period': f"Period {i+1}",
                'start_date': full_portfolio.index[start_idx],
                'end_date': full_portfolio.index[end_idx-1],
                'base_return': base_period_return,
                'full_return': full_period_return,
                'return_diff': full_period_return - base_period_return,
                'avg_sentiment': avg_sentiment,
                'sentiment_volatility': sentiment_volatility
            })
        
        # Compile all results
        sentiment_contribution = {
            'summary': {
                'base_metrics': base_metrics,
                'full_metrics': full_metrics,
                'performance_diff': performance_diff
            },
            'sentiment_driven_trades': sentiment_driven_trades,
            'period_analysis': period_returns
        }
        
        self.logger.info(f"Sentiment contribution analysis complete. Return difference: {performance_diff['return_diff']:.2%}")
        return sentiment_contribution
    
    def generate_sentiment_report(
        self,
        sentiment_analysis: Dict,
        report_path: Optional[str] = None
    ) -> str:
        """
        Generate a report specifically analyzing sentiment contribution.
        
        Args:
            sentiment_analysis: Output from analyze_sentiment_contribution
            report_path: Optional custom path for the report
            
        Returns:
            Path to the generated report
        """
        if report_path is None:
            report_path = os.path.join(self.output_dir, 'sentiment_contribution_report.docx')
        
        doc = Document()
        doc.add_heading('Sentiment Contribution Analysis Report', level=1)
        
        # Add executive summary
        doc.add_heading('Executive Summary', level=2)
        
        summary = sentiment_analysis['summary']
        perf_diff = summary['performance_diff']
        
        summary_text = (
            f"This report analyzes the contribution of Reddit sentiment analysis to trading performance. "
            f"Incorporating sentiment data resulted in a {perf_diff['return_diff']:.2%} "
            f"({'increase' if perf_diff['return_diff'] > 0 else 'decrease'}) in total return, "
            f"which represents a {abs(perf_diff['return_diff_pct']):.2%} {'improvement' if perf_diff['return_diff'] > 0 else 'reduction'} "
            f"over the base strategy. The Sharpe ratio {'improved' if perf_diff['sharpe_diff'] > 0 else 'decreased'} "
            f"by {abs(perf_diff['sharpe_diff']):.2f} points, and maximum drawdown was "
            f"{'reduced' if perf_diff['drawdown_diff'] > 0 else 'increased'} by {abs(perf_diff['drawdown_diff']):.2%}."
        )
        doc.add_paragraph(summary_text)
        
        # Performance comparison table
        doc.add_heading('Performance Comparison', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Base Strategy'
        hdr_cells[2].text = 'With Sentiment'
        
        metrics_to_compare = [
            ('Total Return', f"{summary['base_metrics']['total_return']:.2%}", f"{summary['full_metrics']['total_return']:.2%}"),
            ('Annualized Return', f"{summary['base_metrics']['ann_return']:.2%}", f"{summary['full_metrics']['ann_return']:.2%}"),
            ('Sharpe Ratio', f"{summary['base_metrics']['sharpe_ratio']:.2f}", f"{summary['full_metrics']['sharpe_ratio']:.2f}"),
            ('Max Drawdown', f"{summary['base_metrics']['max_drawdown']:.2%}", f"{summary['full_metrics']['max_drawdown']:.2%}"),
            ('Win Rate', f"{summary['base_metrics']['win_rate']:.2%}", f"{summary['full_metrics']['win_rate']:.2%}"),
            ('Number of Trades', f"{summary['base_metrics']['n_trades']:.0f}", f"{summary['full_metrics']['n_trades']:.0f}")
        ]
        
        for metric, base_val, full_val in metrics_to_compare:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = base_val
            row_cells[2].text = full_val
        
        # Sentiment-driven trades
        doc.add_heading('Sentiment-Driven Trades', level=2)
        trades = sentiment_analysis['sentiment_driven_trades']
        
        if trades:
            doc.add_paragraph(f"The sentiment analysis influenced {len(trades)} trading decisions:")
            
            trade_table = doc.add_table(rows=1, cols=4)
            trade_table.style = 'Table Grid'
            
            hdr_cells = trade_table.rows[0].cells
            hdr_cells[0].text = 'Date'
            hdr_cells[1].text = 'Action'
            hdr_cells[2].text = 'Price'
            hdr_cells[3].text = 'Sentiment'
            
            for trade in trades[:10]:  # Show top 10 trades
                row_cells = trade_table.add_row().cells
                row_cells[0].text = trade['date'].strftime('%Y-%m-%d')
                row_cells[1].text = 'Buy' if trade['order'] > 0 else 'Sell'
                row_cells[2].text = f"${trade['price']:.2f}"
                row_cells[3].text = f"{trade['sentiment']:.4f}" if trade['sentiment'] is not None else 'N/A'
            
            if len(trades) > 10:
                doc.add_paragraph(f"...and {len(trades) - 10} more trades.")
        else:
            doc.add_paragraph("No trades were exclusively driven by sentiment analysis.")
        
        # Period analysis
        doc.add_heading('Market Period Analysis', level=2)
        periods = sentiment_analysis['period_analysis']
        
        period_table = doc.add_table(rows=1, cols=6)
        period_table.style = 'Table Grid'
        
        hdr_cells = period_table.rows[0].cells
        hdr_cells[0].text = 'Period'
        hdr_cells[1].text = 'Date Range'
        hdr_cells[2].text = 'Base Return'
        hdr_cells[3].text = 'With Sentiment'
        hdr_cells[4].text = 'Difference'
        hdr_cells[5].text = 'Avg. Sentiment'
        
        for period in periods:
            row_cells = period_table.add_row().cells
            row_cells[0].text = period['period']
            row_cells[1].text = f"{period['start_date'].strftime('%Y-%m-%d')} to {period['end_date'].strftime('%Y-%m-%d')}"
            row_cells[2].text = f"{period['base_return']:.2%}"
            row_cells[3].text = f"{period['full_return']:.2%}"
            row_cells[4].text = f"{period['return_diff']:.2%}"
            row_cells[5].text = f"{period['avg_sentiment']:.4f}" if period['avg_sentiment'] is not None else 'N/A'
        
        # Add insights
        doc.add_heading('Key Insights', level=2)
        
        # Find period with highest improvement
        best_period = max(periods, key=lambda x: x['return_diff'])
        worst_period = min(periods, key=lambda x: x['return_diff'])
        
        insights = [
            f"Sentiment analysis provided the greatest improvement during {best_period['period']} "
            f"({best_period['start_date'].strftime('%Y-%m-%d')} to {best_period['end_date'].strftime('%Y-%m-%d')}), "
            f"with a {best_period['return_diff']:.2%} increase in returns.",
            
            f"The average sentiment during this period was {best_period['avg_sentiment']:.4f}" 
            if best_period['avg_sentiment'] is not None else "",
            
            f"Sentiment analysis performed worst during {worst_period['period']}, "
            f"with a {worst_period['return_diff']:.2%} impact on returns.",
            
            f"The strategy with sentiment achieved {perf_diff['return_diff_pct']:.2%} higher returns "
            f"than the base strategy without sentiment."
        ]
        
        for insight in insights:
            if insight:  # Only add non-empty insights
                doc.add_paragraph(insight)
        
        # Save the document
        doc.save(report_path)
        self.logger.info(f"Sentiment contribution report saved to {report_path}")
        
        return report_path
    
    def generate_report(
        self,
        portfolio_df: pd.DataFrame,
        metrics: Dict[str, float],
        ticker: str = "Stock",
        report_path: Optional[str] = None
    ) -> str:
        """
        Generate a DOCX report with portfolio performance.
        
        Args:
            portfolio_df: Output from backtest() method
            metrics: Performance metrics dictionary
            ticker: Stock ticker or name for report title
            report_path: Optional custom path for the report
            
        Returns:
            Path to the generated report
        """
        if report_path is None:
            report_path = os.path.join(self.output_dir, 'trading_report.docx')
        
        doc = Document()
        doc.add_heading(f'Trading System Report: {ticker}', level=1)
        
        # Add executive summary
        doc.add_heading('Executive Summary', level=2)
        doc.add_paragraph(f"This report presents the performance of a trading strategy applied to {ticker} "
                         f"using a combination of technical analysis, sentiment analysis, and predictive models.")
                         
        summary_text = (
            f"The strategy achieved a total return of {metrics['total_return']:.2%} "
            f"({metrics['ann_return']:.2%} annualized) with a Sharpe ratio of {metrics['sharpe_ratio']:.2f}. "
            f"Maximum drawdown was {metrics['max_drawdown']:.2%}. "
            f"The strategy executed {metrics['n_trades']:.0f} trades with a win rate of {metrics['win_rate']:.2%}."
        )
        doc.add_paragraph(summary_text)
        
        # Performance Metrics Table
        doc.add_heading('Performance Metrics', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        metrics_to_show = [
            ('Total Return', f"{metrics['total_return']:.2%}"),
            ('Annualized Return', f"{metrics['ann_return']:.2%}"),
            ('Volatility', f"{metrics['volatility']:.2%}"),
            ('Sharpe Ratio', f"{metrics['sharpe_ratio']:.2f}"),
            ('Maximum Drawdown', f"{metrics['max_drawdown']:.2%}"),
            ('Calmar Ratio', f"{metrics['calmar_ratio']:.2f}"),
            ('Win Rate', f"{metrics['win_rate']:.2%}"),
            ('Number of Trades', f"{metrics['n_trades']:.0f}"),
            ('Average Trade Return', f"{metrics['avg_trade']:.2%}")
        ]
        
        for metric, value in metrics_to_show:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        # Equity Curve Plot
        doc.add_heading('Equity Curve', level=2)
        img_stream = io.BytesIO()
        plt.figure(figsize=(8, 4))
        plt.plot(portfolio_df.index, portfolio_df['equity'])
        plt.title('Portfolio Equity Over Time')
        plt.xlabel('Date')
        plt.ylabel('Portfolio Value ($)')
        plt.grid(True)
        plt.tight_layout()
        plt.savefig(img_stream, format='png')
        plt.close()
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(6))
        
        # Drawdown Plot
        doc.add_heading('Drawdown', level=2)
        img_stream = io.BytesIO()
        plt.figure(figsize=(8, 4))
        plt.fill_between(portfolio_df.index, portfolio_df['drawdown'], 0, color='red', alpha=0.3)
        plt.plot(portfolio_df.index, portfolio_df['drawdown'], color='red')
        plt.title('Portfolio Drawdown')
        plt.xlabel('Date')
        plt.ylabel('Drawdown (%)')
        plt.grid(True)
        plt.tight_layout()
        plt.savefig(img_stream, format='png')
        plt.close()
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(6))
        
        # Save the document
        doc.save(report_path)
        self.logger.info(f"Report saved to {report_path}")
        
        return report_path
    
    def create_sentiment_comparison_dashboard(
        self,
        results: Dict,
        config: Dict[str, any]
    ) -> gr.Blocks:
        """
        Create an interactive dashboard for visualizing sentiment contribution.
        
        Args:
            results: Dictionary with base and full strategy results
            config: Configuration parameters
            
        Returns:
            Gradio interface
        """
        # Extract data
        base_portfolio = results['base_portfolio']
        sentiment_portfolio = results['sentiment_portfolio']
        sentiment_df = results.get('sentiment', pd.DataFrame())
        sentiment_analysis = results.get('sentiment_analysis', {})
        
        # Calculate metrics
        base_metrics = self.compute_performance(base_portfolio)
        sentiment_metrics = self.compute_performance(sentiment_portfolio)
        
        # Define dashboard components
        def plot_equity_comparison():
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.plot(base_portfolio.index, base_portfolio['equity'], label='Base Strategy', alpha=0.7)
            ax.plot(sentiment_portfolio.index, sentiment_portfolio['equity'], label='With Sentiment', linewidth=2)
            ax.set_title('Strategy Comparison: Equity Curves')
            ax.set_xlabel('Date')
            ax.set_ylabel('Portfolio Value ($)')
            ax.legend()
            ax.grid(True)
            return fig
        
        def plot_return_comparison():
            # Calculate cumulative returns
            base_cum_returns = base_portfolio['returns'].cumsum()
            sentiment_cum_returns = sentiment_portfolio['returns'].cumsum()
            
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.plot(base_portfolio.index, base_cum_returns, label='Base Strategy', alpha=0.7)
            ax.plot(sentiment_portfolio.index, sentiment_cum_returns, label='With Sentiment', linewidth=2)
            ax.set_title('Strategy Comparison: Cumulative Returns')
            ax.set_xlabel('Date')
            ax.set_ylabel('Cumulative Return')
            ax.legend()
            ax.grid(True)
            return fig
        
        def plot_drawdown_comparison():
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.fill_between(base_portfolio.index, base_portfolio['drawdown'], 0, color='blue', alpha=0.2, label='Base Strategy')
            ax.fill_between(sentiment_portfolio.index, sentiment_portfolio['drawdown'], 0, color='red', alpha=0.2, label='With Sentiment')
            ax.plot(base_portfolio.index, base_portfolio['drawdown'], color='blue', alpha=0.7)
            ax.plot(sentiment_portfolio.index, sentiment_portfolio['drawdown'], color='red')
            ax.set_title('Strategy Comparison: Drawdowns')
            ax.set_xlabel('Date')
            ax.set_ylabel('Drawdown (%)')
            ax.legend()
            ax.grid(True)
            return fig
        
        def plot_sentiment_vs_returns():
            if sentiment_df.empty or 'avg_compound' not in sentiment_df.columns:
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.text(0.5, 0.5, "No sentiment data available", 
                       horizontalalignment='center', verticalalignment='center')
                return fig
            
            # Calculate return difference
            base_returns = base_portfolio['returns']
            sentiment_returns = sentiment_portfolio['returns']
            return_diff = sentiment_returns - base_returns
            
            # Create figure with two Y axes
            fig, ax1 = plt.subplots(figsize=(10, 6))
            
            # Plot return difference on primary Y-axis
            ax1.set_xlabel('Date')
            ax1.set_ylabel('Return Difference', color='tab:blue')
            ax1.plot(return_diff.index, return_diff, color='tab:blue', alpha=0.7)
            ax1.tick_params(axis='y', labelcolor='tab:blue')
            
            # Create secondary Y-axis for sentiment
            ax2 = ax1.twinx()
            ax2.set_ylabel('Sentiment Score', color='tab:red')
            
            # Plot sentiment, handling date alignment
            if isinstance(sentiment_df.index, pd.DatetimeIndex):
                dates = sentiment_df.index
            else:
                dates = pd.to_datetime(sentiment_df['date'])
            
            ax2.plot(dates, sentiment_df['avg_compound'], color='tab:red')
            ax2.tick_params(axis='y', labelcolor='tab:red')
            
            fig.tight_layout()
            fig.suptitle('Sentiment vs. Return Difference', y=1.02)
            
            return fig
        
        def plot_period_analysis():
            if not sentiment_analysis or 'period_analysis' not in sentiment_analysis:
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.text(0.5, 0.5, "Period analysis not available", 
                       horizontalalignment='center', verticalalignment='center')
                return fig
            
            periods = sentiment_analysis['period_analysis']
            
            fig, ax = plt.subplots(figsize=(10, 6))
            
            period_names = [p['period'] for p in periods]
            base_returns = [p['base_return'] * 100 for p in periods]
            sentiment_returns = [p['full_return'] * 100 for p in periods]
            
            x = np.arange(len(period_names))
            width = 0.35
            
            ax.bar(x - width/2, base_returns, width, label='Base Strategy', color='blue', alpha=0.7)
            ax.bar(x + width/2, sentiment_returns, width, label='With Sentiment', color='red', alpha=0.7)
            
            ax.set_ylabel('Period Return (%)')
            ax.set_title('Returns by Market Period')
            ax.set_xticks(x)
            ax.set_xticklabels(period_names)
            ax.legend()
            
            # Add value labels on top of bars
            for i, v in enumerate(base_returns):
                ax.text(i - width/2, v + 0.5, f"{v:.1f}%", ha='center')
            for i, v in enumerate(sentiment_returns):
                ax.text(i + width/2, v + 0.5, f"{v:.1f}%", ha='center')
                
            return fig
        
        def get_metrics_comparison():
            metrics_df = pd.DataFrame({
                'Metric': ['Total Return', 'Annualized Return', 'Sharpe Ratio', 'Max Drawdown', 'Win Rate', 'Trades'],
                'Base Strategy': [
                    f"{base_metrics['total_return'] * 100:.2f}%",
                    f"{base_metrics['ann_return'] * 100:.2f}%",
                    f"{base_metrics['sharpe_ratio']:.2f}",
                    f"{base_metrics['max_drawdown'] * 100:.2f}%",
                    f"{base_metrics['win_rate'] * 100:.2f}%",
                    f"{base_metrics['n_trades']:.0f}"
                ],
                'With Sentiment': [
                    f"{sentiment_metrics['total_return'] * 100:.2f}%",
                    f"{sentiment_metrics['ann_return'] * 100:.2f}%",
                    f"{sentiment_metrics['sharpe_ratio']:.2f}",
                    f"{sentiment_metrics['max_drawdown'] * 100:.2f}%",
                    f"{sentiment_metrics['win_rate'] * 100:.2f}%",
                    f"{sentiment_metrics['n_trades']:.0f}"
                ],
                'Difference': [
                    f"{(sentiment_metrics['total_return'] - base_metrics['total_return']) * 100:.2f}%",
                    f"{(sentiment_metrics['ann_return'] - base_metrics['ann_return']) * 100:.2f}%",
                    f"{sentiment_metrics['sharpe_ratio'] - base_metrics['sharpe_ratio']:.2f}",
                    f"{(base_metrics['max_drawdown'] - sentiment_metrics['max_drawdown']) * 100:.2f}%",
                    f"{(sentiment_metrics['win_rate'] - base_metrics['win_rate']) * 100:.2f}%",
                    f"{sentiment_metrics['n_trades'] - base_metrics['n_trades']:.0f}"
                ]
            })
            return metrics_df
        
        def get_sentiment_driven_trades():
            if not sentiment_analysis or 'sentiment_driven_trades' not in sentiment_analysis:
                return pd.DataFrame({
                    'Trade': ['No sentiment-driven trades data available'],
                    'Value': ['']
                })
            
            trades = sentiment_analysis['sentiment_driven_trades']
            if not trades:
                return pd.DataFrame({
                    'Trade': ['No sentiment-driven trades identified'],
                    'Value': ['']
                })
                
            trade_df = pd.DataFrame(trades)
            # Format dates
            trade_df['date'] = pd.to_datetime(trade_df['date']).dt.strftime('%Y-%m-%d')
            # Add action column
            trade_df['action'] = trade_df['order'].apply(lambda x: 'Buy' if x > 0 else 'Sell')
            
            return trade_df[['date', 'action', 'price', 'sentiment']].rename(
                columns={'date': 'Date', 'action': 'Action', 'price': 'Price', 'sentiment': 'Sentiment'}
            )
        
        # Create Gradio interface
        with gr.Blocks(title="Sentiment Contribution Analysis") as dashboard:
            gr.Markdown(f"# Sentiment Contribution Analysis\n### Technology Stock(s): {', '.join(config.get('tickers', ['Stock']))}")
            
            with gr.Row():
                with gr.Column():
                    gr.Dataframe(get_metrics_comparison(), label="Strategy Performance Comparison")
                
                with gr.Column():
                    gr.Markdown(f"""
                    ## Analysis Parameters
                    - Start Date: {config.get('start_date', 'N/A')}
                    - End Date: {config.get('end_date', 'N/A')}
                    - Initial Capital: ${config.get('initial_capital', 100000.0)}
                    - Reddit Subreddit: {config.get('reddit_subreddit', 'CryptoCurrency')}
                    - Sentiment Weight: {config.get('sentiment_weight', 0.6)}
                    - Technical Weight: {config.get('technical_weight', 0.4)}
                    """)
            
            gr.Markdown("## Equity Curves Comparison")
            gr.Plot(plot_equity_comparison)
            
            with gr.Row():
                with gr.Column():
                    gr.Markdown("## Cumulative Returns")
                    gr.Plot(plot_return_comparison)
                
                with gr.Column():
                    gr.Markdown("## Drawdown Comparison")
                    gr.Plot(plot_drawdown_comparison)
            
            gr.Markdown("## Sentiment vs. Return Difference")
            gr.Plot(plot_sentiment_vs_returns)
            
            gr.Markdown("## Period Analysis")
            gr.Plot(plot_period_analysis)
            
            gr.Markdown("## Sentiment-Driven Trades")
            gr.Dataframe(get_sentiment_driven_trades())
            
        return dashboard
    
    def create_dashboard(
        self,
        results: Dict[str, pd.DataFrame],
        config: Dict[str, any]
    ) -> gr.Blocks:
        """
        Create an interactive dashboard for visualizing results.
        
        Args:
            results: Dictionary with 'portfolio', 'orders', 'metrics', 'sentiment' data
            config: Configuration parameters
            
        Returns:
            Gradio interface
        """
        # Check if we're doing sentiment evaluation
        if 'base_portfolio' in results and 'sentiment_portfolio' in results:
            return self.create_sentiment_comparison_dashboard(results, config)
        
        # Extract data from results
        portfolio_df = results.get('portfolio')
        if portfolio_df is None:
            self.logger.error("No portfolio data found in results")
            return gr.Blocks(title="Error")
            
        metrics = results.get('metrics', {})
        orders_df = results.get('orders', pd.DataFrame())
        sentiment_df = results.get('sentiment', pd.DataFrame())
        
        # Dashboard components
        def plot_equity_curve():
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.plot(portfolio_df.index, portfolio_df['equity'])
            ax.set_title('Equity Curve')
            ax.set_xlabel('Date')
            ax.set_ylabel('Portfolio Value ($)')
            ax.grid(True)
            return fig
        
        def plot_drawdown():
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.fill_between(portfolio_df.index, portfolio_df['drawdown'], 0, color='red', alpha=0.3)
            ax.plot(portfolio_df.index, portfolio_df['drawdown'], color='red')
            ax.set_title('Portfolio Drawdown')
            ax.set_xlabel('Date')
            ax.set_ylabel('Drawdown (%)')
            ax.grid(True)
            return fig
        
        def plot_returns_histogram():
            if 'returns' not in portfolio_df.columns:
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.text(0.5, 0.5, "Returns data not available", 
                     horizontalalignment='center', verticalalignment='center')
                return fig
                
            fig, ax = plt.subplots(figsize=(10, 6))
            portfolio_df['returns'].hist(bins=30, ax=ax)
            ax.set_title('Returns Distribution')
            ax.set_xlabel('Daily Return')
            ax.set_ylabel('Frequency')
            ax.grid(True)
            return fig
        
        def plot_trades():
            if orders_df.empty or 'order' not in orders_df.columns:
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.text(0.5, 0.5, "Trade data not available", 
                     horizontalalignment='center', verticalalignment='center')
                return fig
                
            fig, ax = plt.subplots(figsize=(10, 6))
            # Plot price
            if 'price' in orders_df.columns:
                ax.plot(orders_df.index, orders_df['price'], color='gray', alpha=0.7)
            
            # Plot buy points
            buys = orders_df[orders_df['order'] > 0]
            if not buys.empty and 'price' in buys.columns:
                ax.scatter(buys.index, buys['price'], marker='^', color='green', s=100, label='Buy')
            
            # Plot sell points
            sells = orders_df[orders_df['order'] < 0]
            if not sells.empty and 'price' in sells.columns:
                ax.scatter(sells.index, sells['price'], marker='v', color='red', s=100, label='Sell')
            
            ax.set_title('Trade Signals')
            ax.set_xlabel('Date')
            ax.set_ylabel('Price')
            ax.legend()
            ax.grid(True)
            return fig
        
        def plot_sentiment_vs_price():
            # Check if we have sentiment and price data
            if sentiment_df.empty or 'avg_compound' not in sentiment_df.columns:
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.text(0.5, 0.5, "No sentiment data available", 
                       horizontalalignment='center', verticalalignment='center')
                return fig
            
            if orders_df.empty or 'price' not in orders_df.columns:
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.text(0.5, 0.5, "No price data available", 
                       horizontalalignment='center', verticalalignment='center')
                return fig
            
            # Create figure with two Y axes
            fig, ax1 = plt.subplots(figsize=(10, 6))
            
            # Plot price on primary Y-axis
            ax1.set_xlabel('Date')
            ax1.set_ylabel('Price', color='tab:blue')
            ax1.plot(orders_df.index, orders_df['price'], color='tab:blue')
            ax1.tick_params(axis='y', labelcolor='tab:blue')
            
            # Create secondary Y-axis for sentiment
            ax2 = ax1.twinx()
            ax2.set_ylabel('Sentiment Score', color='tab:red')
            
            # Plot sentiment, handling date alignment
            if isinstance(sentiment_df.index, pd.DatetimeIndex):
                dates = sentiment_df.index
            else:
                dates = pd.to_datetime(sentiment_df['date'])
            
            ax2.plot(dates, sentiment_df['avg_compound'], color='tab:red')
            ax2.tick_params(axis='y', labelcolor='tab:red')
            
            # Add buy/sell points
            buys = orders_df[orders_df['order'] > 0]
            if not buys.empty:
                ax1.scatter(buys.index, buys['price'], marker='^', color='green', s=100, label='Buy')
            
            sells = orders_df[orders_df['order'] < 0]
            if not sells.empty:
                ax1.scatter(sells.index, sells['price'], marker='v', color='red', s=100, label='Sell')
            
            ax1.legend(loc='upper left')
            fig.tight_layout()
            fig.suptitle('Price vs. Reddit Sentiment', y=1.02)
            
            return fig
        
        def get_metrics_table():
            if not metrics:
                return pd.DataFrame({
                    'Metric': ['No metrics data available'],
                    'Value': ['']
                })
                
            metrics_df = pd.DataFrame({
                'Metric': [
                    'Total Return',
                    'Annualized Return',
                    'Sharpe Ratio',
                    'Max Drawdown',
                    'Win Rate',
                    'Number of Trades'
                ],
                'Value': [
                    f"{metrics.get('total_return', 0) * 100:.2f}%",
                    f"{metrics.get('ann_return', 0) * 100:.2f}%",
                    f"{metrics.get('sharpe_ratio', 0):.2f}",
                    f"{metrics.get('max_drawdown', 0) * 100:.2f}%",
                    f"{metrics.get('win_rate', 0) * 100:.2f}%",
                    f"{metrics.get('n_trades', 0):.0f}"
                ]
            })
            return metrics_df
        
        def get_sentiment_stats():
            if sentiment_df.empty or 'avg_compound' not in sentiment_df.columns:
                return pd.DataFrame({
                    'Metric': ['No sentiment data available'],
                    'Value': ['']
                })
            
            sentiment_series = sentiment_df['avg_compound']
            stats = pd.DataFrame({
                'Metric': [
                    'Mean Sentiment',
                    'Median Sentiment',
                    'Max Sentiment',
                    'Min Sentiment',
                    'Sentiment Volatility'
                ],
                'Value': [
                    f"{sentiment_series.mean():.4f}",
                    f"{sentiment_series.median():.4f}",
                    f"{sentiment_series.max():.4f}",
                    f"{sentiment_series.min():.4f}",
                    f"{sentiment_series.std():.4f}"
                ]
            })
            return stats
        
        # Create Gradio interface
        title = "Crypto-Sentiment Trading Dashboard"
        tickers = config.get('tickers', ['Stock'])
        ticker_str = ', '.join(tickers) if tickers else 'Stock'
        
        with gr.Blocks(title=title) as dashboard:
            gr.Markdown(f"# {title}\n### Technology Stock(s): {ticker_str}")
            
            with gr.Row():
                with gr.Column():
                    gr.Dataframe(get_metrics_table(), label="Performance Metrics")
                
                with gr.Column():
                    gr.Markdown(f"""
                    ## Trading Parameters
                    - Start Date: {config.get('start_date', 'N/A')}
                    - End Date: {config.get('end_date', 'N/A')}
                    - Initial Capital: ${config.get('initial_capital', 100000.0)}
                    - Min Holding Period: {config.get('min_holding', 5)} days
                    - Max Trades per Week: {config.get('max_trades', 3)}
                    - Sentiment Method: {config.get('sentiment_method', 'transformer')}
                    - Model Type: {config.get('model_type', 'gru')}
                    - Reddit Subreddit: {config.get('reddit_subreddit', 'CryptoCurrency')}
                    """)
            
            gr.Markdown("## Equity Curve")
            gr.Plot(plot_equity_curve)
            
            gr.Markdown("## Drawdown")
            gr.Plot(plot_drawdown)
            
            # Sentiment and Price Correlation
            gr.Markdown("## Reddit Sentiment vs. Stock Price")
            gr.Plot(plot_sentiment_vs_price)
            
            with gr.Row():
                with gr.Column():
                    gr.Markdown("## Sentiment Statistics")
                    gr.Dataframe(get_sentiment_stats())
                
                with gr.Column():
                    gr.Markdown("## Returns Distribution")
                    gr.Plot(plot_returns_histogram)
            
            gr.Markdown("## Trade Signals")
            gr.Plot(plot_trades)
            
            gr.Markdown("## Recent Trades")
            # Get non-zero trades
            trades = orders_df[orders_df['order'] != 0].tail(10) if 'order' in orders_df.columns else pd.DataFrame()

            # Check if we have any trades to display
            if trades.empty:
                gr.Markdown("No trades were executed in this period.")
            else:
                # Reset index safely
                trades_display = trades.reset_index()
                
                # Check what columns we have available
                available_cols = trades_display.columns.tolist()
                
                # Map expected column names to available ones
                col_mapping = {}
                
                # Handle the index column (typically contains dates)
                if 'index' in available_cols:
                    col_mapping['Date'] = 'index'
                elif trades_display.index.name and trades_display.index.name in available_cols:
                    col_mapping['Date'] = trades_display.index.name
                else:
                    # Use the first column as date if it looks like a date
                    first_col = available_cols[0]
                    if pd.api.types.is_datetime64_any_dtype(trades_display[first_col]):
                        col_mapping['Date'] = first_col
                
                # Handle order and price columns
                if 'order' in available_cols:
                    col_mapping['Order'] = 'order'
                
                if 'price' in available_cols:
                    col_mapping['Price'] = 'price'
                
                # Handle sentiment if available
                if 'sentiment' in available_cols:
                    col_mapping['Sentiment'] = 'sentiment'
                
                # Create display columns list from what's available
                display_cols = []
                for display_name, source_col in col_mapping.items():
                    if source_col in available_cols:
                        display_cols.append(source_col)
                
                if display_cols:
                    # Extract only available columns
                    trades_display = trades_display[display_cols]
                    
                    # Rename to friendly display names
                    rename_map = {source_col: display_name 
                                 for display_name, source_col in col_mapping.items()
                                 if source_col in display_cols}
                    trades_display = trades_display.rename(columns=rename_map)
                    
                    # Format date and add trade type
                    if 'Date' in trades_display.columns:
                        trades_display['Date'] = pd.to_datetime(trades_display['Date']).dt.strftime('%Y-%m-%d')
                    
                    if 'Order' in trades_display.columns:
                        trades_display['Type'] = trades_display['Order'].apply(lambda x: 'Buy' if x > 0 else 'Sell')
                    
                    # Show dataframe
                    final_cols = [col for col in ['Date', 'Type', 'Price', 'Sentiment'] if col in trades_display.columns]
                    
                    if final_cols:
                        gr.Dataframe(trades_display[final_cols])
                    else:
                        gr.Markdown("Trade data available but no displayable columns found.")
                else:
                    gr.Markdown("No displayable trade data available.")
            
        return dashboard


if __name__ == '__main__':
    # Setup basic logging
    logging.basicConfig(level=logging.INFO)
    
    # Test portfolio manager with sample data
    from pandas import date_range
    import matplotlib
    matplotlib.use('Agg')  # Non-interactive backend for testing
    
    # Create sample data
    dates = date_range('2023-01-01', periods=252, freq='B')
    
    # Generate price series (random walk)
    price = 100 + np.cumsum(np.random.randn(len(dates)))
    
    # Generate orders at random intervals
    orders = np.zeros(len(dates))
    for i in range(10, len(dates), 30):  # Every ~30 days
        orders[i] = 1  # Buy
        if i + 15 < len(dates):
            orders[i + 15] = -1  # Sell 15 days later
    
    orders_df = pd.DataFrame({'price': price, 'order': orders}, index=dates)
    
    # Create a sentiment series
    sentiment = np.random.normal(0.2, 0.3, size=len(dates))
    for i in range(1, len(sentiment)):
        sentiment[i] = 0.8 * sentiment[i-1] + 0.2 * sentiment[i]  # Add some autocorrelation
    
    sentiment_df = pd.DataFrame({'avg_compound': sentiment}, index=dates)
    
    # Create portfolio manager
    portfolio_manager = PortfolioManager(initial_capital=100000.0)
    
    # Backtest portfolio
    portfolio_df = portfolio_manager.backtest(orders_df)
    
    # Calculate performance metrics
    metrics = portfolio_manager.compute_performance(portfolio_df)
    print("Performance Metrics:")
    for metric, value in metrics.items():
        print(f"  {metric}: {value}")
    
    # Generate report
    report_path = portfolio_manager.generate_report(
        portfolio_df=portfolio_df,
        metrics=metrics,
        ticker="SAMPLE"
    )
    print(f"Report saved to: {report_path}")
    
    # Test dashboard creation
    results = {
        'portfolio': portfolio_df,
        'orders': orders_df,
        'metrics': metrics,
        'sentiment': sentiment_df
    }
    
    config = {
        'start_date': '2023-01-01',
        'end_date': '2023-12-31',
        'tickers': ['SAMPLE'],
        'initial_capital': 100000.0
    }
    
    dashboard = portfolio_manager.create_dashboard(results, config)
    print("Dashboard created successfully")
    # Note: In a real scenario, you would call dashboard.launch() here