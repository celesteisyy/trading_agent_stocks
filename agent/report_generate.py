import io
import os
import openai
from dotenv import load_dotenv, find_dotenv
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

# Configuration
_ = load_dotenv(find_dotenv("fin580.env"))
api_key = os.getenv('OPENAI_API_KEY')
if not api_key:
    raise ValueError("OPENAI_API_KEY not found in environment. Please set it in your .env file.")
openai.api_key = api_key

class ReportGenerator:
    """
    ReportGenerator creates a DOCX report containing:
      - Title and LLM-generated summary via OpenAI API
      - Performance metrics table
      - Equity curve plot
    """
    def __init__(self, portfolio_df, metrics):
        """
        Args:
            portfolio_df (pd.DataFrame): output from PortfolioManagerAgent.backtest(), must contain 'equity'
            metrics (dict): performance metrics from PortfolioManagerAgent.compute_performance()
        """
        self.portfolio = portfolio_df
        self.metrics = metrics

    def _generate_summary(self) -> str:
        """
        Use OpenAI to generate a text summary based on the performance metrics.
        """
        prompt = (
            "You are a financial analytics assistant. Given these performance metrics, "
            f"provide a concise report summary describing the strategy's performance:\n{self.metrics}"
        )
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=300
        )
        return response.choices[0].message.content.strip()

    def create_report(self, filename: str = 'trading_report.docx'):
        """
        Generate the DOCX report and save under agent/output/<filename>.
        """
        # Generate summary via OpenAI
        summary_text = self._generate_summary()

        doc = Document()
        doc.add_heading('Trading System Report', level=1)
        doc.add_paragraph(summary_text)

        # Performance Metrics Table
        doc.add_heading('Performance Metrics', level=2)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        for metric, value in self.metrics.items():
            row_cells = table.add_row().cells
            row_cells[0].text = metric.replace('_', ' ').title()
            row_cells[1].text = f"{value:.4f}"

        # Equity Curve Plot
        doc.add_heading('Equity Curve', level=2)
        img_stream = io.BytesIO()
        plt.figure(figsize=(6, 3))
        plt.plot(self.portfolio.index, self.portfolio['equity'])
        plt.title('Equity Over Time')
        plt.xlabel('Date')
        plt.ylabel('Portfolio Equity')
        plt.tight_layout()
        plt.savefig(img_stream, format='png')
        plt.close()
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(6))

        # Prepare output path under agent/output
        out_dir = os.path.join('agent', 'output')
        os.makedirs(out_dir, exist_ok=True)
        output_path = os.path.join(out_dir, filename)
        doc.save(output_path)
        print(f"Report saved to {output_path}")

if __name__ == '__main__':
    # Self-check / example usage
    import numpy as np
    import pandas as pd
    from portfolio import PortfolioManagerAgent

    # Generate synthetic price and order data for testing
    dates = pd.date_range('2025-01-01', periods=50, freq='B')
    price = 100 + np.cumsum(np.random.randn(len(dates)))
    df_price = pd.DataFrame({'price': price}, index=dates)

    orders = pd.DataFrame(index=dates)
    orders['order'] = 0
    orders.loc[dates[10], 'order'] = 1
    orders.loc[dates[30], 'order'] = -1
    orders['price'] = df_price['price']

    # Backtest and compute performance
    pm = PortfolioManagerAgent(initial_capital=100000)
    portfolio_df = pm.backtest(orders)
    metrics = pm.compute_performance(portfolio_df)

    # Generate the report
    rg = ReportGenerator(portfolio_df=portfolio_df, metrics=metrics)
    rg.create_report()
